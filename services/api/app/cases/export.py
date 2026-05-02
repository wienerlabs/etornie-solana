"""Case summary export — PDF, Word (DOCX), and Excel (XLSX).

The exported file is a single, self-contained "filing receipt" the user
can hand to their accountant or counsel. It carries every piece of
state that proves the engagement actually happened: the case number,
the on-chain attestation tx, the case NFT mint, the latest filing
robot result, and any x402 payment + compliance proof signatures.

Logo is embedded from ``app/assets/etornie-logo.png`` so the file
renders the same regardless of where the dashboard repo lives.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as PdfImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.cases.models import Case, CaseEvent
from app.documents.models import Document
from app.required_documents.models import CaseRequiredDocument
from app.services.ukipo.models import UKIPOSubmission

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "etornie-logo.png")
LOGO_PATH = os.path.abspath(LOGO_PATH)

ETORNIE_ACCENT = colors.HexColor("#2520FE")


@dataclass(frozen=True)
class ExportContext:
    case: Case
    documents: list[Document]
    required_documents: list[CaseRequiredDocument]
    latest_ukipo: UKIPOSubmission | None
    events: list[CaseEvent]


async def collect_export_context(db: AsyncSession, case: Case) -> ExportContext:
    docs_result = await db.execute(
        select(Document).where(Document.case_id == case.id)
    )
    documents = list(docs_result.scalars().all())

    required_result = await db.execute(
        select(CaseRequiredDocument).where(
            CaseRequiredDocument.case_id == case.id
        )
    )
    required_documents = list(required_result.scalars().all())

    ukipo_result = await db.execute(
        select(UKIPOSubmission)
        .where(UKIPOSubmission.case_id == case.id)
        .order_by(UKIPOSubmission.created_at.desc())
        .limit(1)
    )
    latest_ukipo = ukipo_result.scalar_one_or_none()

    events_result = await db.execute(
        select(CaseEvent)
        .where(CaseEvent.case_id == case.id)
        .order_by(CaseEvent.created_at.asc())
    )
    events = list(events_result.scalars().all())

    return ExportContext(
        case=case,
        documents=documents,
        required_documents=required_documents,
        latest_ukipo=latest_ukipo,
        events=events,
    )


def _format_dt(value: datetime | None) -> str:
    if value is None:
        return "—"
    return value.strftime("%Y-%m-%d %H:%M UTC")


def _format_date(value) -> str:
    if value is None:
        return "—"
    return value.isoformat()


def _nice_classes_to_str(value: str | None) -> str:
    if not value:
        return "—"
    return ", ".join(p.strip() for p in value.split(",") if p.strip())


def _summary_rows(ctx: ExportContext) -> list[tuple[str, str]]:
    case = ctx.case
    return [
        ("Case Number", case.case_number),
        ("Title", case.title or "—"),
        ("Description", case.description or "—"),
        ("Type", case.case_type.value),
        ("Status", case.status.value),
        ("Jurisdiction", case.jurisdiction or "—"),
        ("Nice Classes", _nice_classes_to_str(case.nice_classes)),
        ("Filing Date", _format_date(case.filing_date)),
        ("Deadline", _format_date(case.deadline)),
        ("Client Wallet", case.client_wallet or "—"),
        ("Attestation Tx", case.attestation_tx or "—"),
        ("Attestation PDA", case.attestation_pda or "—"),
        ("NFT State", case.nft_state.value),
        ("NFT Mint", case.nft_mint or "—"),
        ("NFT Mint Tx", case.nft_mint_tx or "—"),
        ("NFT Burn Tx", case.nft_burn_tx or "—"),
        ("Created At", _format_dt(case.created_at)),
        ("Updated At", _format_dt(case.updated_at)),
    ]


def _ukipo_rows(submission: UKIPOSubmission | None) -> list[tuple[str, str]]:
    if submission is None:
        return []
    return [
        ("Submission ID", str(submission.id)),
        ("Status", submission.status.value),
        ("Current Step", submission.current_step or "—"),
        ("Error Step", submission.error_step or "—"),
        ("IPO Reference", submission.ipo_reference or "—"),
        ("IPO Application URL", submission.ipo_application_url or "—"),
        ("Owner", submission.owner_company_name or "—"),
        ("Owner Country", submission.owner_country or "—"),
        ("Mark Type", submission.mark_type.value),
        ("Mark Text", submission.mark_text or "—"),
        ("Payment Tx", submission.solana_payment_tx or "—"),
        (
            "Payment Lamports",
            str(submission.solana_payment_lamports)
            if submission.solana_payment_lamports is not None
            else "—",
        ),
        ("Payment At", _format_dt(submission.solana_payment_at)),
        ("Payer Wallet", submission.solana_payer_wallet or "—"),
        ("Compliance Tx", submission.solana_compliance_tx or "—"),
        ("Compliance PDA", submission.solana_compliance_pda or "—"),
        ("Query Hash", submission.solana_query_hash_hex or "—"),
        ("Commitment", submission.solana_commitment_hex or "—"),
        ("Started At", _format_dt(submission.started_at)),
        ("Finished At", _format_dt(submission.finished_at)),
    ]


def _document_rows(documents: list[Document]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for doc in documents:
        rows.append(
            {
                "filename": doc.filename,
                "type": doc.document_type or "—",
                "status": doc.status.value,
                "ownership_proof_pda": doc.ownership_proof_pda or "—",
                "ownership_verified_at": _format_dt(doc.ownership_verified_at),
                "uploaded_at": _format_dt(doc.created_at),
            }
        )
    return rows


def _required_rows(rows: list[CaseRequiredDocument]) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    for r in rows:
        out.append(
            {
                "document_name": r.document_name,
                "status": r.status.value,
                "linked_document_id": (
                    str(r.document_id) if r.document_id else "—"
                ),
            }
        )
    return out


def _has_logo() -> bool:
    return os.path.isfile(LOGO_PATH)


def render_pdf(ctx: ExportContext) -> bytes:
    """Render the case summary as a one- to two-page A4 PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"Etornie Case {ctx.case.case_number}",
        author="Etornie",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "EtornieTitle",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        textColor=ETORNIE_ACCENT,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "EtornieSubtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.grey,
        spaceAfter=12,
    )
    section_style = ParagraphStyle(
        "EtornieSection",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=ETORNIE_ACCENT,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = styles["BodyText"]

    story: list[Any] = []

    if _has_logo():
        logo = PdfImage(LOGO_PATH, width=22 * mm, height=22 * mm)
        story.append(logo)
        story.append(Spacer(1, 4))

    story.append(Paragraph("Etornie Case Filing Receipt", title_style))
    story.append(
        Paragraph(
            (
                "Generated "
                f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
                " · System of record for on-chain attestation, filing, and proof state."
            ),
            subtitle_style,
        )
    )

    def render_kv_table(rows: list[tuple[str, str]]) -> Table:
        data = [[Paragraph(f"<b>{k}</b>", body_style), Paragraph(v, body_style)] for k, v in rows]
        table = Table(data, colWidths=[55 * mm, None])
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.whitesmoke, colors.white]),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ]
            )
        )
        return table

    story.append(Paragraph("Case Summary", section_style))
    story.append(render_kv_table(_summary_rows(ctx)))

    if ctx.latest_ukipo is not None:
        story.append(Paragraph("Latest UKIPO Submission", section_style))
        story.append(render_kv_table(_ukipo_rows(ctx.latest_ukipo)))

    if ctx.documents:
        story.append(Paragraph("Uploaded Documents", section_style))
        header = ["Filename", "Type", "Status", "Ownership PDA", "Verified At", "Uploaded"]
        rows = [header]
        for d in _document_rows(ctx.documents):
            rows.append(
                [
                    d["filename"],
                    d["type"],
                    d["status"],
                    d["ownership_proof_pda"],
                    d["ownership_verified_at"],
                    d["uploaded_at"],
                ]
            )
        doc_table = Table(rows, repeatRows=1)
        doc_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), ETORNIE_ACCENT),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ]
            )
        )
        story.append(doc_table)

    if ctx.required_documents:
        story.append(Paragraph("Required Documents Checklist", section_style))
        header = ["Document Name", "Status", "Linked Document ID"]
        rows = [header]
        for d in _required_rows(ctx.required_documents):
            rows.append([d["document_name"], d["status"], d["linked_document_id"]])
        req_table = Table(rows, repeatRows=1)
        req_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), ETORNIE_ACCENT),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ]
            )
        )
        story.append(req_table)

    if ctx.events:
        story.append(Paragraph("On-chain Event Timeline", section_style))
        header = ["Event Type", "Tx Signature", "Actor Wallet", "Recorded"]
        rows = [header]
        for ev in ctx.events:
            rows.append(
                [
                    str(ev.event_type),
                    ev.tx_signature,
                    ev.actor_wallet,
                    _format_dt(ev.created_at),
                ]
            )
        ev_table = Table(rows, repeatRows=1)
        ev_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), ETORNIE_ACCENT),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ]
            )
        )
        story.append(ev_table)

    doc.build(story)
    return buffer.getvalue()


def render_docx(ctx: ExportContext) -> bytes:
    """Render the case summary as a Word .docx file."""
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    document = DocxDocument()

    if _has_logo():
        document.add_picture(LOGO_PATH, width=Inches(0.9))

    title = document.add_heading("Etornie Case Filing Receipt", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x25, 0x20, 0xFE)

    subtitle = document.add_paragraph(
        f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} · "
        "System of record for on-chain attestation, filing, and proof state."
    )
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def render_kv_section(heading: str, rows: list[tuple[str, str]]) -> None:
        h = document.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x20, 0xFE)
        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light List Accent 1"
        for i, (k, v) in enumerate(rows):
            cell_k = table.cell(i, 0)
            cell_v = table.cell(i, 1)
            cell_k.text = k
            cell_v.text = v
            for paragraph in cell_k.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    render_kv_section("Case Summary", _summary_rows(ctx))

    if ctx.latest_ukipo is not None:
        render_kv_section(
            "Latest UKIPO Submission",
            _ukipo_rows(ctx.latest_ukipo),
        )

    if ctx.documents:
        h = document.add_heading("Uploaded Documents", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x20, 0xFE)
        table = document.add_table(
            rows=1 + len(ctx.documents),
            cols=6,
        )
        table.style = "Light List Accent 1"
        header_cells = table.rows[0].cells
        for col, label in enumerate(
            ["Filename", "Type", "Status", "Ownership PDA", "Verified At", "Uploaded"]
        ):
            header_cells[col].text = label
        for i, d in enumerate(_document_rows(ctx.documents), start=1):
            row = table.rows[i].cells
            row[0].text = d["filename"]
            row[1].text = d["type"]
            row[2].text = d["status"]
            row[3].text = d["ownership_proof_pda"]
            row[4].text = d["ownership_verified_at"]
            row[5].text = d["uploaded_at"]

    if ctx.required_documents:
        h = document.add_heading("Required Documents Checklist", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x20, 0xFE)
        table = document.add_table(
            rows=1 + len(ctx.required_documents), cols=3
        )
        table.style = "Light List Accent 1"
        header = table.rows[0].cells
        header[0].text = "Document Name"
        header[1].text = "Status"
        header[2].text = "Linked Document ID"
        for i, r in enumerate(_required_rows(ctx.required_documents), start=1):
            row = table.rows[i].cells
            row[0].text = r["document_name"]
            row[1].text = r["status"]
            row[2].text = r["linked_document_id"]

    if ctx.events:
        h = document.add_heading("On-chain Event Timeline", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x20, 0xFE)
        table = document.add_table(rows=1 + len(ctx.events), cols=4)
        table.style = "Light List Accent 1"
        header = table.rows[0].cells
        for col, label in enumerate(
            ["Event Type", "Tx Signature", "Actor Wallet", "Recorded"]
        ):
            header[col].text = label
        for i, ev in enumerate(ctx.events, start=1):
            row = table.rows[i].cells
            row[0].text = str(ev.event_type)
            row[1].text = ev.tx_signature
            row[2].text = ev.actor_wallet
            row[3].text = _format_dt(ev.created_at)

    footer_paragraph = document.add_paragraph()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_paragraph.add_run(
        f"Etornie · Case {ctx.case.case_number}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_xlsx(ctx: ExportContext) -> bytes:
    """Render the case summary as a multi-sheet Excel workbook."""
    workbook = Workbook()
    accent_fill = PatternFill(start_color="2520FE", end_color="2520FE", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    label_font = Font(bold=True)

    summary_sheet = workbook.active
    summary_sheet.title = "Case Summary"
    summary_sheet["A1"] = "Field"
    summary_sheet["B1"] = "Value"
    for col in ("A1", "B1"):
        summary_sheet[col].fill = accent_fill
        summary_sheet[col].font = header_font

    for i, (k, v) in enumerate(_summary_rows(ctx), start=2):
        summary_sheet.cell(row=i, column=1, value=k).font = label_font
        cell = summary_sheet.cell(row=i, column=2, value=v)
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    summary_sheet.column_dimensions[get_column_letter(1)].width = 28
    summary_sheet.column_dimensions[get_column_letter(2)].width = 70

    if ctx.latest_ukipo is not None:
        sheet = workbook.create_sheet("UKIPO Submission")
        sheet["A1"] = "Field"
        sheet["B1"] = "Value"
        for col in ("A1", "B1"):
            sheet[col].fill = accent_fill
            sheet[col].font = header_font
        for i, (k, v) in enumerate(_ukipo_rows(ctx.latest_ukipo), start=2):
            sheet.cell(row=i, column=1, value=k).font = label_font
            cell = sheet.cell(row=i, column=2, value=v)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        sheet.column_dimensions[get_column_letter(1)].width = 28
        sheet.column_dimensions[get_column_letter(2)].width = 70

    if ctx.documents:
        sheet = workbook.create_sheet("Documents")
        headers = [
            "Filename",
            "Type",
            "Status",
            "Ownership PDA",
            "Verified At",
            "Uploaded",
        ]
        for col, label in enumerate(headers, start=1):
            cell = sheet.cell(row=1, column=col, value=label)
            cell.fill = accent_fill
            cell.font = header_font
        for i, d in enumerate(_document_rows(ctx.documents), start=2):
            sheet.cell(row=i, column=1, value=d["filename"])
            sheet.cell(row=i, column=2, value=d["type"])
            sheet.cell(row=i, column=3, value=d["status"])
            sheet.cell(row=i, column=4, value=d["ownership_proof_pda"])
            sheet.cell(row=i, column=5, value=d["ownership_verified_at"])
            sheet.cell(row=i, column=6, value=d["uploaded_at"])
        for col in range(1, len(headers) + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 30

    if ctx.required_documents:
        sheet = workbook.create_sheet("Required Documents")
        headers = ["Document Name", "Status", "Linked Document ID"]
        for col, label in enumerate(headers, start=1):
            cell = sheet.cell(row=1, column=col, value=label)
            cell.fill = accent_fill
            cell.font = header_font
        for i, r in enumerate(_required_rows(ctx.required_documents), start=2):
            sheet.cell(row=i, column=1, value=r["document_name"])
            sheet.cell(row=i, column=2, value=r["status"])
            sheet.cell(row=i, column=3, value=r["linked_document_id"])
        for col in range(1, len(headers) + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 36

    if ctx.events:
        sheet = workbook.create_sheet("Event Timeline")
        headers = ["Event Type", "Tx Signature", "Actor Wallet", "Recorded"]
        for col, label in enumerate(headers, start=1):
            cell = sheet.cell(row=1, column=col, value=label)
            cell.fill = accent_fill
            cell.font = header_font
        for i, ev in enumerate(ctx.events, start=2):
            sheet.cell(row=i, column=1, value=str(ev.event_type))
            sheet.cell(row=i, column=2, value=ev.tx_signature)
            sheet.cell(row=i, column=3, value=ev.actor_wallet)
            sheet.cell(row=i, column=4, value=_format_dt(ev.created_at))
        for col in range(1, len(headers) + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 36

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
