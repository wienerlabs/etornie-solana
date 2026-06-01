"""Render a GDPR data export (the dict from :mod:`data_export`) to a
branded PDF, Word (DOCX), or Excel (XLSX) document.

JSON stays the canonical, machine-readable Article-20 format; these
renderers are the human-readable companions a user can hand to counsel
or an accountant. They consume the very same ``build_user_export`` dict,
so the three formats can never drift from the JSON.

The export is a nested structure — a metadata header, a ``profile``
key/value object, and a list of row-objects per domain table — so the
renderers are generic: every section is drawn from the dict's shape
rather than from hand-listed fields, which keeps them correct as new
tables are added to the export.
"""
from __future__ import annotations

import io
import json
import os
from typing import Any, Iterator
from xml.sax.saxutils import escape as _xml_escape

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as PdfImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

LOGO_PATH = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "assets", "etornie-logo.png")
)
ETORNIE_ACCENT_HEX = "2520FE"
_ETORNIE_ACCENT = colors.HexColor(f"#{ETORNIE_ACCENT_HEX}")

# Section keys whose Title-cased form would read wrong.
_SECTION_LABELS = {
    "etorniegpt_chat_messages": "EtornieGPT Chat Messages",
    "in_app_notifications": "In-App Notifications",
    "braid_feedback_events": "BRAID Feedback Events",
}

_NULL = "—"


def _label(key: str) -> str:
    return _SECTION_LABELS.get(key, key.replace("_", " ").title())


def _cell(value: Any) -> str:
    """Render one value as a flat string cell."""
    if value is None:
        return _NULL
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (str, int, float)):
        return str(value)
    # lists / dicts (e.g. JSONB columns) round-trip as compact JSON.
    return json.dumps(value, ensure_ascii=False)


def _meta_rows(export: dict[str, Any]) -> list[tuple[str, str]]:
    subject = export["subject"]
    return [
        ("Export format", _cell(export["export_format"])),
        ("Export version", _cell(export["export_version"])),
        ("GDPR basis", _cell(export["gdpr_basis"])),
        ("Generated at", _cell(export["generated_at"])),
        ("User ID", _cell(subject.get("user_id"))),
        ("Email", _cell(subject.get("email"))),
        ("Full name", _cell(subject.get("full_name"))),
        ("Wallet address", _cell(subject.get("wallet_address"))),
        ("Public handle", _cell(subject.get("public_handle"))),
    ]


def _profile_rows(profile: dict[str, Any]) -> list[tuple[str, str]]:
    return [(_label(k), _cell(v)) for k, v in profile.items()]


def _collection_sections(
    export: dict[str, Any],
) -> Iterator[tuple[str, list[dict[str, Any]]]]:
    """Yield ``(label, rows)`` for every list-valued data section."""
    for key, value in export["data"].items():
        if isinstance(value, list):
            yield _label(key), value


def _union_keys(rows: list[dict[str, Any]]) -> list[str]:
    """Column order = first-seen order across every row."""
    keys: list[str] = []
    seen: set[str] = set()
    for row in rows:
        for key in row:
            if key not in seen:
                seen.add(key)
                keys.append(key)
    return keys


def _counts(export: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        (_label(key), str(len(value)))
        for key, value in export["data"].items()
        if isinstance(value, list)
    ]


def _has_logo() -> bool:
    return os.path.isfile(LOGO_PATH)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def render_pdf(export: dict[str, Any]) -> bytes:
    """Render the export as an A4 PDF — one key/value block per record."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="Etornie Personal Data Export",
        author="Etornie",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title_", parent=styles["Title"], fontSize=18, leading=22,
        textColor=_ETORNIE_ACCENT, spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle_", parent=styles["Normal"], fontSize=9,
        textColor=colors.grey, spaceAfter=12,
    )
    section_style = ParagraphStyle(
        "Section_", parent=styles["Heading2"], fontSize=12,
        textColor=_ETORNIE_ACCENT, spaceBefore=12, spaceAfter=6,
    )
    item_style = ParagraphStyle(
        "Item_", parent=styles["Heading3"], fontSize=9,
        textColor=colors.grey, spaceBefore=8, spaceAfter=2,
    )
    body = styles["BodyText"]

    def kv_table(rows: list[tuple[str, str]]) -> Table:
        # reportlab Paragraph parses a mini-HTML grammar, so any '<', '>'
        # or '&' in the real data (chat text, JSON values, URLs) must be
        # XML-escaped or the build raises a parse error.
        data = [
            [
                Paragraph(f"<b>{_xml_escape(k)}</b>", body),
                Paragraph(_xml_escape(v), body),
            ]
            for k, v in rows
        ]
        table = Table(data, colWidths=[55 * mm, None])
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1),
                     [colors.whitesmoke, colors.white]),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ]
            )
        )
        return table

    story: list[Any] = []
    if _has_logo():
        story.append(PdfImage(LOGO_PATH, width=22 * mm, height=22 * mm))
        story.append(Spacer(1, 4))
    story.append(Paragraph("Personal Data Export", title_style))
    story.append(
        Paragraph(
            _xml_escape(
                f"Generated {export['generated_at']} · {export['gdpr_basis']}"
            ),
            subtitle_style,
        )
    )

    story.append(Paragraph("Export Details", section_style))
    story.append(kv_table(_meta_rows(export)))
    story.append(Paragraph("Record Counts", section_style))
    story.append(kv_table(_counts(export)))

    story.append(Paragraph("Profile", section_style))
    story.append(kv_table(_profile_rows(export["data"]["profile"])))

    # Collection records are rendered as flowable key/value paragraphs
    # rather than table rows: a single record can carry a JSONB blob
    # (tool results, gateway metadata) taller than a page, and a table
    # row cannot split across pages whereas a paragraph can.
    for label, rows in _collection_sections(export):
        story.append(Paragraph(f"{label} ({len(rows)})", section_style))
        if not rows:
            story.append(Paragraph("No records.", body))
            continue
        for index, row in enumerate(rows, start=1):
            story.append(Paragraph(f"#{index}", item_style))
            for key, value in row.items():
                story.append(
                    Paragraph(
                        f"<b>{_xml_escape(_label(key))}:</b> "
                        f"{_xml_escape(_cell(value))}",
                        body,
                    )
                )

    doc.build(story)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_docx(export: dict[str, Any]) -> bytes:
    """Render the export as a Word document — a section per data table."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor

    accent = RGBColor(0x25, 0x20, 0xFE)
    grey = RGBColor(0x6B, 0x72, 0x80)
    document = DocxDocument()

    if _has_logo():
        document.add_picture(LOGO_PATH, width=Inches(0.9))

    title = document.add_heading("Personal Data Export", level=0)
    for run in title.runs:
        run.font.color.rgb = accent
    subtitle = document.add_paragraph(
        f"Generated {export['generated_at']} · {export['gdpr_basis']}"
    )
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = grey

    def kv_section(heading: str, rows: list[tuple[str, str]]) -> None:
        head = document.add_heading(heading, level=1)
        for run in head.runs:
            run.font.color.rgb = accent
        if not rows:
            document.add_paragraph("No records.")
            return
        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light List Accent 1"
        for i, (key, value) in enumerate(rows):
            cell_k, cell_v = table.cell(i, 0), table.cell(i, 1)
            cell_k.text = key
            cell_v.text = value
            for paragraph in cell_k.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    kv_section("Export Details", _meta_rows(export))
    kv_section("Record Counts", _counts(export))
    kv_section("Profile", _profile_rows(export["data"]["profile"]))

    for label, rows in _collection_sections(export):
        head = document.add_heading(f"{label} ({len(rows)})", level=1)
        for run in head.runs:
            run.font.color.rgb = accent
        if not rows:
            document.add_paragraph("No records.")
            continue
        for index, row in enumerate(rows, start=1):
            sub = document.add_heading(f"#{index}", level=2)
            for run in sub.runs:
                run.font.color.rgb = grey
            table = document.add_table(rows=len(row), cols=2)
            table.style = "Light List Accent 1"
            for i, (key, value) in enumerate(row.items()):
                cell_k, cell_v = table.cell(i, 0), table.cell(i, 1)
                cell_k.text = _label(key)
                cell_v.text = _cell(value)
                for paragraph in cell_k.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def _sanitize_sheet_title(title: str, used: set[str]) -> str:
    """Excel sheet titles: ≤31 chars, none of ``[]:*?/\\``, unique."""
    cleaned = title
    for bad in "[]:*?/\\":
        cleaned = cleaned.replace(bad, " ")
    cleaned = cleaned.strip()[:31] or "Sheet"
    candidate = cleaned
    suffix = 2
    while candidate.lower() in used:
        tail = f" {suffix}"
        candidate = cleaned[: 31 - len(tail)] + tail
        suffix += 1
    used.add(candidate.lower())
    return candidate


def render_xlsx(export: dict[str, Any]) -> bytes:
    """Render the export as a workbook — one sheet per data table."""
    workbook = Workbook()
    accent_fill = PatternFill(
        start_color=ETORNIE_ACCENT_HEX,
        end_color=ETORNIE_ACCENT_HEX,
        fill_type="solid",
    )
    header_font = Font(color="FFFFFF", bold=True)
    label_font = Font(bold=True)
    used_titles: set[str] = set()

    def kv_sheet(title: str, rows: list[tuple[str, str]]) -> None:
        sheet = workbook.create_sheet(_sanitize_sheet_title(title, used_titles))
        sheet["A1"], sheet["B1"] = "Field", "Value"
        for col in ("A1", "B1"):
            sheet[col].fill = accent_fill
            sheet[col].font = header_font
        for i, (key, value) in enumerate(rows, start=2):
            sheet.cell(row=i, column=1, value=key).font = label_font
            sheet.cell(row=i, column=2, value=value).alignment = Alignment(
                wrap_text=True, vertical="top"
            )
        sheet.column_dimensions[get_column_letter(1)].width = 30
        sheet.column_dimensions[get_column_letter(2)].width = 80

    # First (default) sheet holds the export metadata + record counts.
    overview = workbook.active
    overview.title = "Export Details"
    overview["A1"], overview["B1"] = "Field", "Value"
    for col in ("A1", "B1"):
        overview[col].fill = accent_fill
        overview[col].font = header_font
    row_cursor = 2
    for key, value in _meta_rows(export) + [("", "")] + _counts(export):
        overview.cell(row=row_cursor, column=1, value=key).font = label_font
        overview.cell(row=row_cursor, column=2, value=value)
        row_cursor += 1
    overview.column_dimensions[get_column_letter(1)].width = 30
    overview.column_dimensions[get_column_letter(2)].width = 80

    kv_sheet("Profile", _profile_rows(export["data"]["profile"]))

    for label, rows in _collection_sections(export):
        sheet = workbook.create_sheet(_sanitize_sheet_title(label, used_titles))
        if not rows:
            sheet["A1"] = "No records."
            continue
        columns = _union_keys(rows)
        for col, key in enumerate(columns, start=1):
            cell = sheet.cell(row=1, column=col, value=_label(key))
            cell.fill = accent_fill
            cell.font = header_font
        for i, row in enumerate(rows, start=2):
            for col, key in enumerate(columns, start=1):
                sheet.cell(row=i, column=col, value=_cell(row.get(key)))
        for col in range(1, len(columns) + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 28

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
